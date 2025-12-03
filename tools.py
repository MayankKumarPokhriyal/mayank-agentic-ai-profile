from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

import pandas as pd
from fpdf import FPDF
from docx import Document

# -----------------------------
# Profile Loader
# -----------------------------

_PROFILE_CACHE: Optional[Dict[str, Any]] = None
_PROFILE_PATH = Path(__file__).resolve().parent / "profile.json"
LEADS_DIR = Path(__file__).resolve().parent / "leads"
LEADS_DIR.mkdir(exist_ok=True)


def _load_profile() -> Dict[str, Any]:
    global _PROFILE_CACHE
    if _PROFILE_CACHE is None:
        with _PROFILE_PATH.open("r", encoding="utf-8") as f:
            _PROFILE_CACHE = json.load(f)
    return _PROFILE_CACHE


def get_profile_section(section_name: str):
    if not section_name:
        return {}

    profile = _load_profile()
    section_name = section_name.lower().strip()

    # ✅ Allow common aliases
    aliases = {
        "skill": "skills",
        "skills": "skills",
        "education": "education",
        "experience": "experience",
        "projects": "projects",
        "job": "job_preferences",
        "preferences": "job_preferences",
        "contact": "contact",
        "links": "links"
    }

    key = aliases.get(section_name, section_name)

    return profile.get(key, {})


def get_project_details(project_name: str):
    projects = _load_profile().get("projects", [])
    for project in projects:
        if project["name"].lower() == project_name.lower():
            return project
    return None


# -----------------------------
# Recruiter Lead Export (CSV, PDF, DOCX)
# -----------------------------

def log_recruiter_lead(lead: Dict[str, Any]) -> Dict[str, Any]:
    timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")

    lead_row = {
        "timestamp": timestamp,
        "recruiter_name": lead.get("recruiter_name"),
        "company": lead.get("company"),
        "role": lead.get("role"),
        "contact": lead.get("contact"),
        "notes": lead.get("notes"),
    }

    df = pd.DataFrame([lead_row])

    # ---- CSV Export ----
    csv_path = LEADS_DIR / "recruiter_leads.csv"
    if csv_path.exists():
        df.to_csv(csv_path, mode="a", header=False, index=False)
    else:
        df.to_csv(csv_path, index=False)

    # ---- PDF Export ----
    pdf_path = LEADS_DIR / f"lead_{timestamp.replace(':','-')}.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for k, v in lead_row.items():
        pdf.cell(0, 10, f"{k.upper()}: {v}", ln=True)

    pdf.output(str(pdf_path))

    # ---- DOCX Export ----
    docx_path = LEADS_DIR / f"lead_{timestamp.replace(':','-')}.docx"
    doc = Document()
    doc.add_heading("Recruiter Lead", level=1)
    for k, v in lead_row.items():
        doc.add_paragraph(f"{k.upper()}: {v}")
    doc.save(docx_path)

    return {
        "status": "saved_locally",
        "csv": str(csv_path),
        "pdf": str(pdf_path),
        "docx": str(docx_path),
        "timestamp": timestamp,
    }